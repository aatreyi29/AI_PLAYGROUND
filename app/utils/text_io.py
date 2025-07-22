import os
import sys
import re
import json
import base64
import hashlib
import logging
import requests
import concurrent.futures
from io import BytesIO
from datetime import datetime, timedelta, timezone
from concurrent.futures import ThreadPoolExecutor

from dotenv import load_dotenv
from tqdm import tqdm
from PIL import Image
from docx import Document
import PyPDF2
import fitz  # PyMuPDF
import easyocr
from extract_msg import Message

from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.prompts import ChatPromptTemplate
from langchain.memory import ConversationBufferWindowMemory
from langchain_core.output_parsers import StrOutputParser, JsonOutputParser
from langchain_core.runnables import RunnablePassthrough, RunnableLambda
from langchain_core.messages import HumanMessage, AIMessage
from langchain_community.vectorstores import FAISS
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_aws import ChatBedrockConverse
from openai import OpenAI

from core.logger import logging
from core.exception import CustomException

load_dotenv()

## -------------------------------- playground ------------------------------

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

api_key = os.getenv("OPENAI_API_KEY")

chat_history = ChatMessageHistory()
chat_history.add_user_message("hi!")

persist_directory="DB"
embeddings = OpenAIEmbeddings()
vectordb = FAISS.load_local("./faiss_index", embeddings, allow_dangerous_deserialization=True)

memory = ConversationBufferWindowMemory(k=51,memory_key="chat_history", return_messages=True,)
llm = ChatOpenAI(model="gpt-3.5-turbo")

def generate_session_id(document):
    hasher = hashlib.sha256()
    hasher.update(document.read())
    document.seek(0)  
    hash_bytes = hasher.digest()
    hash_base64 = base64.b64encode(hash_bytes).decode('utf-8')
    session_id = re.sub(r'[^a-zA-Z]', '', hash_base64)
    return session_id[:32]  

# def save_uploaded_image(uploaded_file, output_folder, image_name="uploaded_image"):
#     os.makedirs(output_folder, exist_ok=True)
#     try:
#         image = Image.open(uploaded_file)
#         image_path = os.path.join(output_folder, f"{image_name}.png")
#         image.save(image_path, format="PNG")
#         return image_path
#     except Exception as e:
#         raise ValueError(f"Error processing image file: {e}")

def encode_image(image_path):
    with open(image_path, "rb") as image_file:
        print(f"This is image Path : {image_path}")
        return base64.b64encode(image_file.read()).decode('utf-8')

def process_image(image_path):
    base64_image = encode_image(image_path)
    llm = ChatOpenAI(model="gpt-4")

    prompt = ChatPromptTemplate.from_messages([
        HumanMessage(content=[
            {
                "type": "text",
                "text": "Extract all the text from the image and provide it in a JSON format with a key 'extracted_text'. Ensure you don't break the sequence of text in the image."
            },
            {
                "type": "image_url",
                "image_url": f"data:image/jpeg;base64,{base64_image}"
            }
        ])
    ])

    json_parser = JsonOutputParser()
    chain = prompt | llm | json_parser

    response = chain.invoke({})
    print(f"This is extracted text : {response.get('extracted_text', '')}")

    return response.get('extracted_text', '')


def process_image_1(image_path):
    base64_image = encode_image(image_path)
    client = OpenAI()
    
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": "What is in this image? can you give me all content and text available in image.",
                },
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"},
                },
            ],
        }
    ],
    )
    return response.choices[0].message.content

def convert_pdf_to_images(pdf_path, output_folder, image_format='jpg', quality=85):
    """
    Converts each page of a PDF file to an image.

    Parameters:
        pdf_path (str): Path to the PDF file.
        output_folder (str): Folder where the images will be saved.
        image_format (str): Format of the output images ('jpg' or 'png').
        quality (int): Quality of the output images (1-100, applicable for JPEG).

    Returns:
        list: List of paths to the saved images.
    """
    # Ensure output folder exists
    os.makedirs(output_folder, exist_ok=True)
    
    pdf_document = fitz.open(pdf_path)
    image_paths = []

    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        pix = page.get_pixmap()

        image_path = os.path.join(output_folder, f"page_{page_num + 1}.{image_format}")

        if image_format.lower() == 'jpg':
            pix.save(image_path)  
            # Optionally, use PIL for quality control
            from PIL import Image
            with Image.open(image_path) as img:
                img.save(image_path, format='JPEG', quality=quality)
        elif image_format.lower() == 'png':
            pix.save(image_path)

        image_paths.append(image_path)

    return image_paths

# def process_images_concurrently(image_paths):
#     texts = []
#     with ThreadPoolExecutor() as executor:
#         futures = {executor.submit(process_image, image_path): image_path for image_path in image_paths}
#         for future in tqdm(concurrent.futures.as_completed(futures), total=len(futures), desc="Processing images"):
#             try:
#                 text = future.result()
#                 texts.append(text)
#             except OSError as e:
#                 logger.error(f"OSError processing {futures[future]}: {e}")
#             except Exception as e:
#                 logger.error(f"OSError processing {futures[future]}: {e}")
#     return texts

# def process_image_pdf(doc_path):
#     output_folder = "output_images"
#     if not os.path.exists(output_folder):
#         os.makedirs(output_folder)
#     image_paths = convert_pdf_to_images(doc_path, output_folder)
#     print(f"Image paths : {image_paths}")

#     final_response = ""
#     for path_ in image_paths:
#         processed_response =  process_image_1(path_)
#         print(f"This is processed response : {processed_response}")
#         final_response += processed_response
#     return final_response
        
# def process_doc(doc_path):
#     return process_image_pdf(doc_path)

def extract_text_from_pdf(pdf_file):
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    count = 0
    null_count = 0
    for page_num in range(len(pdf_reader.pages)):
        count += 1
        page = pdf_reader.pages[page_num]
        page_text = page.extract_text()
        if page_text or page_text != "":
            text += page_text
        else:
            null_count += 1

    return text, count, null_count

def extract_text_from_msg(msg_file_path):
    try:
        msg = Message(msg_file_path)
        msg_sender = msg.sender
        msg_date = msg.date
        msg_subject = msg.subject
        msg_body = msg.body

        attachments = msg.attachments

        full_text = f"From: {msg_sender}\nDate: {msg_date}\nSubject: {msg_subject}\n\n{msg_body}"

        for attachment in attachments:
            attachment_data = attachment.data
            attachment_name = attachment.longFilename

            # Example: Save attachment to a file
            # with open(attachment_name, 'wb') as f:
            #     f.write(attachment_data)

        return full_text

    except Exception as e:
        print(f"Error extracting text from .msg file: {e}")
        return None

# def extract_tables_from_docx(file_path):
#     doc = Document(file_path)
#     parsed_tables = []

#     for table in doc.tables:
#         # Extract headers from the first row
#         headers = [cell.text.strip() for cell in table.rows[0].cells]
#         rows = []
#         for row in table.rows[1:]:
#             row_data = [cell.text.strip() for cell in row.cells]
#             rows.append(row_data)

#         parsed_tables.append({'headers': headers, 'rows': rows})

#     return parsed_tables

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = []

    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    # Initialize easyocr Reader
    reader = easyocr.Reader(['en'])

    # Extract text from images
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image = Image.open(BytesIO(image_data))
            image_text = reader.readtext(image, detail=0)
            full_text.append(' '.join(image_text))

    return '\n'.join(full_text)

# Noted
def extract_key_value_pairs(json_obj, parent_key='', sep='_'):
    items = {}
    for k, v in json_obj.items():
        new_key = f"{parent_key}{sep}{k}" if parent_key else k
        if isinstance(v, dict):
            items.update(extract_key_value_pairs(v, new_key, sep=sep))
        elif isinstance(v, list):
            for i, item in enumerate(v):
                items.update(extract_key_value_pairs({f"{new_key}{sep}{i}": item}, '', sep=sep))
        else:
            items[new_key] = v
    return item

## -------------------------------- playground ------------------------------

## -------------------------------- text analyser ------------------------------

# def get_w2_data(user_file):
#     logging.info("Nanonets api called")
#     try:
#         url = 'https://app.nanonets.com/api/v2/OCR/Model/995d3217-671e-4747-befa-c171ea5d5dac/LabelUrls/?async=false'
#         headers = {'accept': 'application/x-www-form-urlencoded'}
#         data = {'file': user_file}
#         response = requests.post(url, headers=headers, auth=requests.auth.HTTPBasicAuth('f185aded-08a3-11ef-a49a-d642b16313d7', ''), files=data)
#         return response.text
#     except Exception as e:
#         logging.info(f"Exception occure during nanonets api call {e}")
#         raise CustomException(e, sys)

# Format documents
# def format_docs(docs):
#     return "\n\n".join([d.page_content for d in docs])

## ------------------------------------------ Conversational_Chain ------------------------------------

# def get_conversational_chain():
#     try:
#         template = """
#             You are an exceptional tax-advisory chatbot, meticulously tailored to provide me with highly personalized guidance based on my W-2 form data. I am expecting extremely detailed suggestions, but easy enough to process for someone without a tax or financial background. I do not want generic advice like 'consult a tax professional'. Treat me as a beginner, but empower me to become a master of my taxes. Use the data you are provided to respond with personalised answers. Answer to the best of your abilities. For any questions that are out of your scope, you can simply tell the user you do not have the capability of answering. Do not just spew the user's W2 data back to them. Use numbers where you feel they are necessary. Respond with bullet points or numbered lists to make it easier for the common person to understand. Do not make any text bold. Answer in a professional and conversational tone. Use the latest tax data from 2024 as reference, but use 2023 if 2024 data is not available.

#             Here some relevant strategies and points of reference that can be used to improve your answers: {context}

#             Here is the user's w2 form data you are expected to base your calculations off: {data}

#             Here are the latest tax limits you are expected to remember: {limits}

#             Here is the chat history you are expected to remember: {memory}

#             Here is the report that has been generated, regarding which the user can ask questions: {recommendations}

#             Here is the user's question: {question}
#             """
#         prompt = ChatPromptTemplate.from_template(template)
#         llm = ChatOpenAI(model="gpt-3.5-turbo")
#         chain = (
#                 RunnablePassthrough.assign(memory=RunnableLambda(memory.load_memory_variables))
#                 )| prompt | llm | StrOutputParser()
#         return chain
#     except Exception as e:
#         logging.info(f"Exception occurs during making chat conversation chain {e}")
#         raise CustomException(e, sys)

## ------------------------------------------ text analyser ------------------------------------

def savings_chain():
    try:
        template = """
        Your objective is to provide detailed, specific tax savings guidance to the tax filer. Base your suggestions on the filer’s W-2 Boxes 1 through 14, and provide strategies of how they can save current and future taxes.  Make responses easy enough to process for someone without a tax or financial background.  In your analysis, please try to include the following items when reasonable: 1) 401(k)/403(b)/457 - pre-tax vs Roth, 2) IRA - traditional, Roth, non-deductible with Roth conversion, 3) HSA (health savings account), 4) FSA (flexible savings account), 5) dependent care benefits, and 6) emergency savings withdrawal of $1,000. If the filer is already taking advantage of certain strategies such as 401(k), HSA, etc…please mention something along the lines of “Great job for taking advantage of…”.  For the IRA recommendation, please consider income phaseout limits based on a Single taxpayer, and if they make too much for a Roth IRA please mention the backdoor Roth option. Do not give abstract and arbitrary suggestions like 'consult a tax or investment professional'. Do not include each and every data point from the dataset in your answer, but use specific numbers where necessary. Your response should ONLY be bullet points, without any headers, introduction or conclusion. Each bullet point should have a prefix of '&&&', and should not be in bold. Only provide your greatest 6-7 points, not more. Do not forget these rules.

        You also need to add required information from context for create more better response.

        Here is the filer's W2 data: {w2_data}

        Here is the context : {context}
        """
        prompt = ChatPromptTemplate.from_template(template)
        llm = ChatOpenAI(model="gpt-3.5-turbo")
        savings_chain = prompt | llm | StrOutputParser()
        return savings_chain
    except Exception as e:
        logging.info(f"Exception occure during creation of saving chain {e}")
        raise CustomException(e, sys)


def withholding_chain():
    try:
        template = """
        Your objective is to provide brief information regarding the tax filer’s Federal income tax withholding rate.  Make responses easy enough to process for someone without a tax or financial background.  Please give them a percentage with no decimal points of their Box 2 amount divided by Box 1, which tells them their Federal tax withholding rate.  Then compare their Federal tax withholding to the IRS tax tables based on a Single filer, and give a basic recommendation of they are likely under withholding or over with holding Federal income tax.  Please tell them this analysis is based on filing as Single, and to use the result with utmost caution.  Also tell them to reach out to their tax advisor to more accurately determine if they are likely under or over withholding based on their more recent paystub, and considering their other income sources.

        You also need to add required information from context for create more better response.

        Here is the filer's W2 data: {w2_data}

        Here is the context : {context}
        """
        prompt = ChatPromptTemplate.from_template(template)
        llm = ChatOpenAI(model="gpt-3.5-turbo")
        investment_chain = prompt | llm | StrOutputParser()
        return investment_chain
    except Exception as e:
        logging.info(f"Exception occure during create investment chain {e}")
        raise CustomException(e, sys)


def tax_credits_chain():
    try:
        template = """
        Your objective is to provide detailed guidance on claiming the following tax credits: 1) Child Tax Credit, 2) Child and Dependent Care Credit, 3) American Opportunity Tax Credit, 4) Lifetime Learning Credit, and 5) Residential Energy Credits.  Make responses easy enough to process for someone without a tax or financial background.  Your response should ONLY be bullet points, without any headers, introduction or conclusion. Each bullet point should have a prefix of '&&&', and should not be in bold. Only provide your greatest 5-6 points, not more. Do not forget these rules.

        You also need to add required information from context for create more better response.

        Here is the filer's W2 data: {w2_data}

        Here is the context : {context}
        """
        prompt = ChatPromptTemplate.from_template(template)
        llm = ChatOpenAI(model="gpt-3.5-turbo")
        optimization_chain = prompt | llm | StrOutputParser()
        return optimization_chain
    except Exception as e:
        logging.info(f"Exception occure during creation optimization chain")
        raise CustomException(e, sys)
    
## -------------------------------------------- Conversational_Chains ------------------------------------

# Noted 
def conversational_chain_revised(question, text): 
    template = f"""
    You are an expert in every global law domain, and your job is to do the best you can to help those who come to you for help. You must respond to queries professionally, ensuring that you do not make mistakes.You will be working with those who will have less technical domain knowledge, and hence you will have to explain concepts clearly and succinctly.

    Here is the content on which you need to answer : {text}

    Here is the user's question: {question}
    """
    prompt = ChatPromptTemplate.from_template(template)
    llm=ChatOpenAI(model="gpt-4o")
    chain = prompt | llm | StrOutputParser()

    response = chain.invoke({"question": question, "text": text})
    return response

# Noted
def conversational_chain_revised_json(question, text):
    template = """
    Follow the given instructions in the user's question as best you can. Return JSON, with no errors in formatting (no introduction or conclusion, only a valid JSON object). Use the given content and text as the raw data.

    Here is the user's question: {question}

    Here is the content on which you need to answer: {text}

    """
    prompt = ChatPromptTemplate.from_template(template)
    llm=ChatOpenAI(model="gpt-4o")
    print("LLM :", llm)
    chain = prompt | llm | JsonOutputParser()

    response = chain.invoke({"question": question, "text": text})
    print(response)
    return response

# def conversational_chain(question, chat_history):
#     template = f"""
#     You are an expert in every global law domain, and your job is to do the best you can to help those who come to you for help. You must respond to queries professionally, ensuring that you do not make mistakes. You will be working with those who will have less technical domain knowledge, and hence you will have to explain concepts clearly and succinctly.

#     Here is the chat history you are expected to remember: {chat_history}

#     Here is the user's question: {question}
#     """
#     prompt = ChatPromptTemplate.from_template(template)
#     llm=ChatOpenAI(model="gpt-4o")
#     chain = prompt | llm | StrOutputParser()

#     response = chain.invoke({"question": question, "chat_history": chat_history})
#     return response

# # Set up conversational chain - legal analyser
# def chat_chain(question, chat_history):
#     template = f"""
#     You are an expert in every global law domain, and your job is to do the best you can to those who come to you for help. You must respond to queries professionally, ensuring that you do not make mistakes.You will be working with those who will have less technical domain knowledge, and hence you will have to explain concepts clearly and succinctly.

#     Remember below suggestions:
#     1. You have to behave like legal adviser.
#     2. Response should be given by only bases of history do not make any answer from your self.
#     3. Answer should be easy to understand and proper related to user question and given content as history.
#     4. If somebody ask any question about out legal content or given history content then tell them to ask questions about given content.

#     Here is the chat history you are expected to remember: {chat_history}

#     Here is the user's question: {question}
#     """
#     prompt = ChatPromptTemplate.from_template(template)
#     llm=ChatOpenAI(model="gpt-4o")
#     chain = prompt | llm | StrOutputParser()

#     response = chain.invoke({"question": question, "chat_history": chat_history})
#     return response

## -------------------------------------------- Conversational_Chain ------------------------------------

## -------------------------------------------- Legal Prompt --------------------------------------------

# prompe_template_legal
prompt_template = """
You are an expert in global law, and you must answer questions as best you can. Understand the information provided below well, as all questions will be related to it. Answer all relevant questions, especially about the document (context). Answer professionally, citing sources, clauses and numbers where you can. Be clear and concise, and do not guess.

Use this information below as reference:

Chat history: \n{chat_history}\n

Context:\n {context}?\n
Question: \n{question}\n

Answer:
"""

# template = """
# You are an expert in every global law domain, and your job is to do the best you can to help those who come to you for help. You must respond to queries professionally, ensuring that you do not make mistakes. Understand the user's question extremely well, and answer it very precisely. You will be working with those who will have less technical domain knowledge, and hence you will have to explain concepts clearly and succinctly.

# Here is the content on basis you need to answer : {text}

# Here is the user's question: {question}
# """

summary_input = """
    Using the context you have been provided, generate a succint and accurate summary of the entire document. Do NOT provide a generic response, such as the document contains XYZ. Be incredibly specific, outline the key terms, use numbers if they are provided, and uncover the essentials. Ensure that you explain what kind of document it is, who are the parties outlined in the document, what is its purpose and what is the gist of it. By reading your response, the reader should have a very clear idea of what this document contains. Your summary should be 5-7 sentences.
    """

outline_input = """
    Using the context you have been provided, produce a crisp, to the point outline of the key clauses the parties need to be wary of. Ensure that you outline the implications provided in the document, as well as any noteworthy conditions. Remember, the user is relying on you to provide a crisp analysis.

    Remember these suggestions when answering:

    1. Structure your output as bullet points, with each bullet focused on a unique clause outline its clause number clearly if there is one mentioned.
    2. Add the '&&&' symbol after each bullet point. Do not forget this. Ensure there are 3 '&' symbols, no more and no less. Ensure each bullet point is prefixed with only one bullet, and no numbers.
    3. Total number of words of overall response should not exceeed 175 words.
    4. Do not show the number of words on screen.
    5. Do not make any of the text bold. Ensure no '*' asterisks are present in your response.

    """

action_items_input = """
    Using the context you have been provided, identify the key action items and deadlines written in the document and make the user aware of it. Remember to be extremely specific about each item, do not be vague. Don't describe the action item, but provide all the information necessary to make good use of that information.

    Remember these suggestions when answering:

    1. Structure your output as bullet points, with each bullet focused on a unique action item.
    2. Add the '&&&' symbol after each bullet point. Do not forget this. Ensure there are 3 '&' symbols, no more and no less. Ensure each bullet point is prefixed with only one bullet, and no numbers.
    3. Total number of words of overall response should not exceeed 175 words.
    4. Do not show the number of words on screen.
    5. Do not make any of the text bold. Ensure no '*' asterisks are present in your response.
    6. Do not give your response a header.

    """

red_flags_input = """
    Using the context you have been provided, carefully search for potential red flags in the document. Explain what the reader must be careful about. Do not provide false information. Be extremely precise when explaining each red flag. Do not provide generic advice.

    Remember these suggestions when answering:

    1. Structure your output as bullet points, with each bullet focused on a unique red flag.
    2. Add the '&&&' symbol after each bullet point. Do not forget this. Ensure there are 3 '&' symbols, no more and no less. Ensure each bullet point is prefixed with only one bullet, and no numbers.
    3. Total number of words of overall response should not exceeed 175 words.
    4. Do not show the number of words on screen.
    5. Do not make any of the text bold. Ensure no '*' asterisks are present in your response.
    6. Do not give your response a header.

    """

outline_money_input = """
    Using the context you have been provided, carefully search for mentions of financial terms. Be extremely precise and succinct in mentioning it. Do not fabricate any numbers, but be sure to involve all financial terms proposed in the document. Ensure each of the bullet point focuses on a numerical terms proposed in the document. If a bullet point does not correspond to a number, do not add it. Every single financial figure, in any currency, should be present in your analysis.

    Remember these suggestions when answering:

    1. Structure your output as bullet points, with each bullet focused on a unique financial figure.
    2. Add the '&&&' symbol after each bullet point. Do not forget this. Ensure there are 3 '&' symbols, no more and no less. Ensure each bullet point is prefixed with only one bullet, and no numbers.
    3. Total number of words of overall response should not exceeed 175 words.
    4. Do not show the number of words on screen.
    5. Do not make any of the text bold. Ensure no '*' asterisks are present in your response.
    6. Do not give your response a header.
    6. Do not provide any header, such as 'Here are financial figures'.

    """

write_questions_input = """
    Using the context you have been provided, carefully understand the entire document, and propose questions the reader can ask their attorney to make best use of the document. Outline beginner friendly questions, but do not add your reasoning for it. Select the top 5 questions and only add those to the list. Ensure they are high quality questions.

    Remember these suggestions when answering:

    1. Structure your output as bullet points, with each bullet focused on a unique question.
    2. Saperate each bullet point of output by '&&&' symbol. Do not forget this. Ensure there are 3 '&' symbols, no more and no less. Do not prefix the bullet points with numbers, only a bullet point.
    3. Total number of words of overall response should not exceeed 175 words.
    4. Do not show the number of words on screen.
    5. Do not make any of the text bold.
    6. Do not add any header.
    """

explain_glossary_input = """
    Using the context you have been provided, carefully search for key legal terms and jargon, and provide an extremely simple explanation of what that term means, and what context it is used in, in this document. Each term should be a unique term. Be sure to keep your explanation short, crisp and helpful. Be sure to include at least 3-5 terms.

    Remember these suggestions when answering:

    1. Structure your output as bullet points, with each bullet focused on a unique term.
    2. Add the '&&&' symbol after each bullet point. Do not forget this. Ensure there are 3 '&' symbols, no more and no less. Ensure each bullet point is prefixed with only one bullet, and no numbers.
    3. Total number of words of overall response should not exceeed 175 words.
    4. Do not show the number of words on screen.
    5. Do not make any of the text bold. Ensure no '*' asterisks are present in your response.
    6. Do not give your response a header.

    """

## -------------------------------------------- Legal Prompt --------------------------------------------

## -------------------------------------------- Policy Prompt --------------------------------------------

prompt_template_policy = """
You are an expert in Global ESG standards, and you must answer questions as best you can. Understand the information provided below well, as all questions will be related to it. Answer all relevant questions, especially about the document (context). Answer professionally, citing sources, clauses and numbers where you can. Be clear and concise, and do not guess.

Use this information below as reference:

Chat history: \n{chat_history}\n

Context:\n {context}?\n
Question: \n{question}\n

Answer:
"""

summary_input_policy = """
    Using the context you have been provided, generate a succint and accurate summary of the entire document. Do NOT provide a generic response, such as the document contains XYZ. Be incredibly specific, outline the key terms, use numbers if they are provided, and uncover the essentials. Ensure you outline the overarching theme of the policy act, and uncover its prominent trends. By reading your response, the reader should have a very clear idea of what this document contains. Your summary should be 7-9 sentences.
    """

outline_input_policy = """
    Using the context you have been provided, produce a crisp, to the point outline of the key clauses in the document. Each clause should pertain to a specific figure or statement present in the clause, and include a brief conclusion of what that figure implicates. Remember, the user is relying on you to provide a crisp analysis.

    Remember these suggestions when answering:

    1. Structure your output as bullet points, with each bullet focused on a unique clause outline its clause number clearly if there is one mentioned.
    2. Add the '&&&' symbol after each bullet point. Do not forget this. Ensure there are 3 '&' symbols, no more and no less. Ensure each bullet point is prefixed with only one bullet, and no numbers.
    3. Do not show the number of words on screen.
    4. Do not make any of the text bold. Ensure no '*' asterisks are present in your response.

    """

intended_audience_input_policy = """
    Using the context you have been provided, identify all the parties that will be affected by these Policy changes. Identify each of the profiles, such as public companies, partnerships, etc. and mention industry if applicable. Each audience should be clearly mentioned in the policy. Do not make guesses or simply provide just the number. Explain how audiences may be affected clearly by each figure.

    Remember these suggestions when answering:

    1. Structure your output as bullet points, with each bullet focused on a unique action item.
    2. Add the '&&&' symbol after each bullet point. Do not forget this. Ensure there are 3 '&' symbols, no more and no less. Ensure each bullet point is prefixed with only one bullet, and no numbers.
    3. Do not show the number of words on screen.
    4. Do not make any of the text bold. Ensure no '*' asterisks are present in your response.
    5. Do not give your response a header.
    6. Do not make any of your text bold.
    """

action_items_input_policy = """
    Using the context you have been provided, identify the key action items and deadlines written in the document and make the user aware of it. Remember to be extremely specific about each item, do not be vague. Cite the numbers and figures present in the document well. Ensure each action item is relevant, not simply an abstract number. Provide the context of each action item clearly and thoughtfully, highlighting why that number is relevant. Ultimately, each point should include what the user needs to look out for (depending on their profile) with this new policy change.

    Remember these suggestions when answering:

    1. Structure your output as bullet points, with each bullet focused on a unique action item.
    2. Add the '&&&' symbol after each bullet point. Do not forget this. Ensure there are 3 '&' symbols, no more and no less. Ensure each bullet point is prefixed with only one bullet, and no numbers.
    3. Do not show the number of words on screen.
    4. Do not make any of the text bold. Ensure no '*' asterisks are present in your response.
    5. Do not give your response a header.
    6. Do not make any of your text bold.
    """

## -------------------------------------------- Policy Prompt --------------------------------------------

## -------------------------------------------- Extraction Prompt --------------------------------------------

json_output = """
    You are a world class data extractor you are master in extract all data and structure them in proper way. your task is to extract all the details from given data. The data you are getting is application form data you need to give each and every detail from the data whether it is checkbox data or anything else. You need to provide every small details from the data of an application form or anything else. If data have box type data or checkbox type data then you need to extract that data. Your task is to convert this data into a structured JSON format. Ensure that you cover all mandatory fields mentioned in the raw data and organize them appropriately.

    Note : Do not break sequence. You need to give response in proper sequense as it is in data. You do not miss any data from given data you must have to give all data in json format no matter if it is json structurized or not you need to make key for that data and present as json format.
"""

# checkbox_output = """
#     You are a world class data extractor you are master in extract data and structure them in proper way. You need to give response in sequence from given data. your task is to extract all the checkbox data from given data. You need to give Quetion/Title of checkbox. Your task is to convert this data into a structured JSON format.Ensure that you need extract only checkbox data do not extract any other details from data.

    # N : Do not break sequence. You need to give response in proper sequence as it is in given data.
# """

# student_detail = """
#     You need to extract students all personal and financial details from given data and need to return in proper json format we any one can easily convert it into table.

#     You need to extract as below format 
    
#     {
#         "Personal_details" : {
#             "Name" : "Name of student", 
#             "Age" : "Age of student",
#             etc ...                     ## You need to attach all othere details like wise in same format
#         } ,
#         "Financial_details" : {
#             "Fees" : "paid fees", 
#             "due" : "due fees", 
#             etc ...                    ## You need to attach all othere details like wise in same format
#         } 
#     }
# """

# Noted
student_verification = """
    You need to verify data from given student data and given text data that all student data is matches in available text if there is any mismatch or descrepancy in data then you need to give it in json format as given below.

    {
        "Descrepancies" :
        {
            "Title of descrepancy" : "Data which has descrepancy along with reason.", 
            "Title of descrepancy" : "Data which has descrepancy along with reason.", 
            etc ...     ## Likewise you need to give all descrepancies here if you found.

        }
    }
    
    Note : You need to check whether name is same in all given text , whether financial detail is matching or not. so this kind of descrepancies you need to provide.
"""

student_sign_verification = """
You need to extract and structure all the data into a proper JSON format as given below, based on the text. 
Ensure that both student signatures and initials are identified accurately, even if initials appear as part of the name (e.g., "Student's initials").

Expected JSON Format:

    {
            "Page_no": "number of page",
            "Initials_sign_available": "Yes/no "  ## You need to check whether Initials sign available or not in that page,
            "No_initials_sign_required": "number response in intiger"  ## No of places where initial sign required on page,
            "No_of_initials_sign_available": "number response in intiger" ## Yu need to number out of all required initial signs how many places sign done.
            "Student_sign_available": "Yes/no "  ## You need to check whether Student sign available or not in that page,
            "No_student_sign_required": "number response in intiger"  ## No of places where Student sign required on page,
            "No_of_student_sign_available": "number response in intiger" ## You need to number out of all required Students signs how many places sign done.
        }

Key Instructions:
1. **Student_sign_available**: Indicate 'Yes' if a student signature is present, otherwise 'No'. but do not consider other sign like witness sign or anything else.
2. **Initials_sign_available**: Indicate 'Yes' if student initials are present, even if specified as "Student's initials". but do not consider other sign like witness sign or anything else.
3. **No_student_sign_required**: Give all number of student signature required where student needs to sign. but do not consider other sign like witness sign or anything else.
4. **No_initials_sign_required**: Give all number of initial signature required where Initial needs to sign. but do not consider other sign like witness sign or anything else.
5. **No_of_initials_sign_available**: Out of all initial signs how many signs did in current page.
6. **No_of_student_sign_available**: Out of all student signs how many signs did in current page.

also you can add another signs if available in page just add in json structure with this 3 key values.
follow below example : 


"Witness_sign_available" : value
"No_witness_sign_required" : value
"No_of_witness_sign_available" : value

if you find any other signature like this please add as above in given json structure below all 7 key values.

Note : You must nee to give output in given json structure and also must need to extract all information from given image.

"""

# Noted
student_sign_verification_structurize = """
You need to extract and structure all the data into a proper JSON format as given below, based on the text. 

Expected JSON Format:

    [{
            "Page_no": "1",
            "Initials_sign_available": "Yes/no "  ## You need to check whether Initials sign available or not in that page,
            "No_initials_sign_required": "number response in intiger"  ## No of places where initial sign required on page,
            "No_of_initials_sign_available": "number response in intiger" ## Yu need to number out of all required initial signs how many places sign done.
            "Student_sign_available": "Yes/no "  ## You need to check whether Student sign available or not in that page,if there is any other signature instead of student signature don't take it as an output
            "No_student_sign_required": "number response in intiger"  ## No of places where Student sign required on page,
            "No_of_student_sign_available": "number response in intiger" ## Yu need to number out of all required Students signs how many places sign done.if there is any other signature instead of student signature don't take it as an output
        }, 
        {
            "Page_no": "2",
            "Initials_sign_available": "Yes/no "  ## You need to check whether Initials sign available or not in that page,
            "No_initials_sign_required": "number response in intiger"  ## No of places where initial sign required on page,
            "No_of_initials_sign_available": "number response in intiger" ## Yu need to number out of all required initial signs how many places sign done.
            "Student_sign_available": "Yes/no "  ## You need to check whether Student sign available or not in that page,if there is any other signature instead of student signature don't take it as an output,
            "No_student_sign_required": "number response in intiger"  ## No of places where Student sign required on page,
            "No_of_student_sign_available": "number response in intiger" ## Yu need to number out of all required Students signs how many places sign done.if there is any other signature instead of student signature don't take it as an output
        },
        {
            "Page_no": "3",
            "Initials_sign_available": "Yes/no "  ## You need to check whether Initials sign available or not in that page,
            "No_initials_sign_required": "number response in intiger"  ## No of places where initial sign required on page,
            "No_of_initials_sign_available": "number response in intiger" ## Yu need to number out of all required initial signs how many places sign done.
            "Student_sign_available": "Yes/no "  ## You need to check whether Student sign available or not in that page, if there is any other signature instead of student signature don't take it as an output,
            "No_student_sign_required": "number response in intiger"  ## No of places where Student sign required on page,
            "No_of_student_sign_available": "number response in intiger" ## Yu need to number out of all required Students signs how many places sign done, if there is any other signature instead of student signature don't take it as an output,.
        }, 
        {
    "Page_no": No of page,
    "Initials_sign_available": "answer",
    "No_initials_sign_required": answer,
    "No_of_initials_sign_available": answer,
    "Student_sign_available": "answer",
    "No_student_sign_required": answer,
    "No_of_student_sign_available": answer,
    "FAA_sign_available": "answer",
    "No_FAA_sign_required": answer,
    "No_of_FAA_sign_available": "answer"
}
]

Note : You must need to give output in proper json structure without miss any field or data.

"""

extract_checkboxes = """

You are a data extraction model. Given the following image, extract all required checkbox fields in below suggested format :

Format your response like this:

{
Question: [Extracted Question]
Options:
- A: [Option A]
- B: [Option B]
- C: [Option C]
- D: [Option D]
Answer: [Correct Option]  ## If none of the option tick then return "None" if all options are ticked then return all available option in list.
}

Note : It is must that you extract all checkbox data from given image.

"""

# Noted
def conversational_chain_student(question, student_detail, text):
    template = """
    Follow the given instructions in the user's question as best you can. Return JSON, with no errors in formatting (no introduction or conclusion, only a valid JSON object). Use the given content and text as the raw data.

    Here is the user's question: {question}
    
    Here is the student Detail : {student_detail}

    Here is the content on which you need to answer: {text}

    """
    prompt = ChatPromptTemplate.from_template(template)
    # llm = ChatBedrockConverse(model="us.meta.llama3-2-11b-instruct-v1:0", max_tokens=None, temperature=0)
    llm=ChatOpenAI(model="gpt-4o")
    print("LLM :", llm)
    chain = prompt | llm | JsonOutputParser()

    response = chain.invoke({"question": question, "student_detail": student_detail, "text": text})
    print(response)
    return response

# Noted
def structure_student_verification(question, text):
    template = """
    Follow the given instructions in the user's question as best you can. Return JSON, with no errors in formatting (no introduction or conclusion, only a valid JSON object). Use the given content and text as the raw data.

    Remember do not miss any data given in text add all data in given format.
    
    
    Here is the user's question: {question}

    Here is the content on which you need to answer: {text}

    """
    prompt = ChatPromptTemplate.from_template(template)
    # llm = ChatBedrockConverse(model="us.meta.llama3-2-11b-instruct-v1:0", max_tokens=None, temperature=0)
    llm=ChatOpenAI(model="gpt-4o")
    print("LLM :", llm)
    chain = prompt | llm | JsonOutputParser()

    response = chain.invoke({"question": question, "text": text})
    print(response)
    return response


def process_image_3(image_path):
    
    base64_image = encode_image(image_path)
    client = OpenAI()
    
    response = client.chat.completions.create(
    model="gpt-4o",
    messages=[
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": student_sign_verification,
                },
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"},
                },
            ],
        }
    ],
)
    return response.choices[0].message.content

# Noted
def process_image_pdf_3(doc_path):
    output_folder = "output_images"
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    image_paths = convert_pdf_to_images(doc_path, output_folder)
    print(f"Image paths : {image_paths}")
    # processed_response =  process_image(image_paths)
    # print(f"This is processed response : {processed_response}")
    # return processed_response
    final_response = ""
    # final_response_json = []
    for path_ in image_paths:
        processed_response =  process_image_3(path_)
        print(f"This is processed response : {processed_response}")
        # final_json_response = json.loads(processed_response)
        # final_response_json.append(final_json_response)
        final_response += processed_response
    # return final_response_json, final_response
    return final_response

def process_image_4(image_path):
    
    base64_image = encode_image(image_path)
    client = OpenAI()
    
    response = client.chat.completions.create(
    model="gpt-4o",
    messages=[
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": extract_checkboxes,
                },
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"},
                },
            ],
        }
    ],
)
    return response.choices[0].message.content

# Noted
def process_image_pdf_4(doc_path):
    output_folder = "output_images"
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    image_paths = convert_pdf_to_images(doc_path, output_folder)
    print(f"Image paths : {image_paths}")
    # processed_response =  process_image(image_paths)
    # print(f"This is processed response : {processed_response}")
    # return processed_response
    final_response = ""
    # final_response_json = []
    for path_ in image_paths:
        processed_response =  process_image_4(path_)
        print(f"This is processed response : {processed_response}")
        # final_json_response = json.loads(processed_response)
        # final_response_json.append(final_json_response)
        final_response += processed_response
    # return final_response_json, final_response
    return final_response

# Noted
structure_student_response_prompt = """  

you need to structurize given text in below format.

[{
Question: [Extracted Question]
Options:
- A: [Option A]
- B: [Option B]
- C: [Option C]
- D: [Option D]
Answer: [Correct Option]
}, 
{
Question: [Extracted Question]
Options:
- A: [Option A]
- B: [Option B]
- C: [Option C]
- D: [Option D]
Answer: [Correct Option]
}, 
{
Question: [Extracted Question]
Options:
- A: [Option A]
- B: [Option B]
- C: [Option C]
- D: [Option D]
Answer: [Correct Option]
}
]

Note : Do not miss any data from given text take all data and return final list of jsons with all extracted question and answer fields.

"""

# Noted
def structure_checkbox_response(question, text):
    template = """
    Follow the given instructions in the user's question as best you can. Return JSON, with no errors in formatting (no introduction or conclusion, only a valid JSON object). Use the given content and text as the raw data.

    Remember do not miss any data given in text add all data in given format.


    Here is the user's question: {question}

    Here is the content on which you need to answer: {text}

    """
    prompt = ChatPromptTemplate.from_template(template)
    # llm = ChatBedrockConverse(model="us.meta.llama3-2-11b-instruct-v1:0", max_tokens=None, temperature=0)
    llm=ChatOpenAI(model="gpt-4o")
    print("LLM :", llm)
    chain = prompt | llm | JsonOutputParser()

    response = chain.invoke({"question": question, "text": text})
    print(response)
    return response


## -------------------------------------------- Extraction Prompt --------------------------------------------

## -------------------------------------------- Medical Prompt --------------------------------------------

json_output_medical = """
    Extract key values and associated information from the given document, focusing on specific sections. The OCR system should capture the following details: Key Values, Measurement Units, Actual Measurements, Ranges, whether the Actual Measurements are within range (show 'true') or out of range (show 'false'). Additionally, provide any structured data found in the document, such as tables, and present the full extracted data in a structured JSON format. Do not include the patient's name and the lab's name; only show the test name and report.

    NOTE : Do not miss any single line of data give all the from given data except patient's name and the lab's name. Do not change key name key name will be TestName, ActualMeasurement, MeasurementUnit, ReferenceRange, IsValueNormal .


    json
    {
  "ReportName": "HAEMATOLOGY INVESTIGATIONS",
  "Tests": [
      {
      "TestName": "Complete Blood Count",
      "TestResults": "Normal",
      "ExpectedRange": "4.5-11.0 x10^3/µL",
      "ActualValue": "7.2 x10^3/µL",
      "IsValueNormal": true
      },
      {
      "TestName": "Blood Glucose",
      "TestResults": "High",
      "ExpectedRange": "70-99 mg/dL",
      "ActualValue": "105 mg/dL",
      "IsValueNormal": false
      }
  ]
  }
"""

personal_output_medical = """
    You need to extract only personal details from given data like Name, Gender, Age. And also remember do not change key name. You need to give extracted data in a structured JSON format.

    Note : Do not give any other information you need to give only Name, Gender, Age.
    {
        'Name': 'patient name',
        'Gender': 'patient gender',
        'Age': 'patient age'
    }
"""

## -------------------------------------------- Medical Prompt --------------------------------------------

## -------------------------------------------- Tax Prompt --------------------------------------------

prompt_template_tax = """
You are a highly knowledgeable and reliable tax advisor specialized in U.S. personal income tax. Your goal is to assist the taxpayer in understanding and optimizing their federal tax situation. Use the W-2 form data provided below as the sole basis for your response. Be specific, practical, and easy to understand, especially for users without a financial background.

Focus on answering accurately and professionally in the following areas when applicable:
- Tax Savings: Offer personalized tips such as contributing to 401(k), IRA, HSA, FSA, etc.
- Tax Withholding: Calculate withholding rate and compare with IRS recommendations.
- Tax Credits: Inform about potential credits like Child Tax Credit, Education Credits, Energy Credits, etc.

Always cite specific boxes or amounts from the W-2 data when appropriate. Do **not** guess. Do **not** provide disclaimers like "consult a professional." Be brief and focused. Bullet points are encouraged.

Reference Material:
Chat History:
{chat_history}

W-2 Context:
{context}

Question:
{question}

Answer:
"""

## -------------------------------------------- Tax Prompt --------------------------------------------